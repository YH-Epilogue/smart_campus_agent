"""
Smart Campus Agent 项目文档生成器
生成包含功能说明和使用指南的 Word 文档
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ========== 样式设置 ==========
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

# ========== 封面 ==========
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Smart Campus Agent')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 100, 200)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('校园智能问答与事务处理平台')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('项目功能说明书及使用指南')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(120, 120, 120)

doc.add_page_break()

# ========== 目录 ==========
toc_title = doc.add_heading('目录', level=1)

toc_items = [
    '一、项目概述',
    '二、技术架构',
    '三、功能模块详解',
    '    3.1 用户认证与权限管理',
    '    3.2 知识库管理',
    '    3.3 文档管理',
    '    3.4 AI 对话系统',
    '    3.5 RAG 检索引擎',
    '    3.6 工具调用系统',
    '    3.7 日志与数据分析',
    '    3.8 系统配置',
    '    3.9 多模态输入',
    '四、部署指南',
    '    4.1 环境要求',
    '    4.2 后端部署',
    '    4.3 前端部署',
    '    4.4 一键启动脚本',
    '五、使用指南',
    '    5.1 登录系统',
    '    5.2 知识库操作',
    '    5.3 文档上传与管理',
    '    5.4 智能对话',
    '    5.5 工具调用示例',
    '    5.6 日志查看与导出',
    '六、API 接口说明',
    '七、常见问题',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ========== 一、项目概述 ==========
doc.add_heading('一、项目概述', level=1)

doc.add_paragraph(
    'Smart Campus Agent 是一款基于 RAG（检索增强生成）技术的校园智能问答与事务处理平台。'
    '该系统整合学生手册、校园新闻、规章制度等校园知识，通过 AI 对话为师生提供信息查询和事务办理服务。'
)

doc.add_paragraph('系统核心能力包括：')

features = [
    '知识库问答 —— 上传校园文档后，AI 可基于文档内容回答问题',
    '多轮对话 —— 支持上下文追问，AI 记住对话历史',
    '工具调用 —— 支持学生信息查询、请假申请、请假状态查询',
    '多模态输入 —— 支持上传图片进行 OCR 文字识别',
    '权限管理 —— 基于角色的访问控制（管理员、知识库管理员、普通用户）',
    '日志分析 —— 对话记录、数据统计、CSV 导出',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()

# ========== 二、技术架构 ==========
doc.add_heading('二、技术架构', level=1)

doc.add_heading('2.1 技术栈', level=2)

table = doc.add_table(rows=1, cols=3, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = '层级'
hdr[1].text = '技术'
hdr[2].text = '说明'

tech_data = [
    ('前端', 'Vue 3 + Vite', '响应式 SPA，Element Plus 组件库'),
    ('前端状态', 'Pinia', '会话管理、用户状态'),
    ('后端框架', 'FastAPI', '异步高性能 API 服务'),
    ('数据库', 'SQLite', '轻量级关系数据库'),
    ('向量数据库', 'ChromaDB', '文档向量存储与检索'),
    ('大语言模型', 'DeepSeek API', '对话生成与理解'),
    ('Embedding', 'sentence-transformers', '文本向量化（text2vec-base-chinese）'),
    ('OCR', 'PaddleOCR', '图片文字识别'),
    ('缓存', 'Redis / 内存', '接口响应缓存，自动降级'),
    ('密码加密', 'PBKDF2-SHA256', '10 万轮迭代 + 随机盐值'),
    ('认证', 'JWT', 'Token 无状态认证'),
]
for layer, tech, desc in tech_data:
    row = table.add_row().cells
    row[0].text = layer
    row[1].text = tech
    row[2].text = desc

doc.add_paragraph()

doc.add_heading('2.2 项目结构', level=2)

structure = (
    'smart_campus_agent/\n'
    '├── backend/                    # FastAPI 后端\n'
    '│   ├── app/\n'
    '│   │   ├── api/v1/            # 7 个 API 路由模块（34 个端点）\n'
    '│   │   ├── services/          # 6 个核心服务（RAG/LLM/NLU/规则/工具/多模态）\n'
    '│   │   ├── models/            # 数据库模型、Pydantic Schema\n'
    '│   │   ├── core/              # 配置、安全、缓存\n'
    '│   │   └── main.py            # FastAPI 入口\n'
    '│   ├── data/                  # 数据目录\n'
    '│   ├── .env                   # 环境变量\n'
    '│   └── requirements.txt       # Python 依赖\n'
    '├── frontend/                   # Vue 3 前端\n'
    '│   ├── src/\n'
    '│   │   ├── views/             # 5 个页面组件\n'
    '│   │   ├── components/        # 5 个可复用组件\n'
    '│   │   ├── stores/            # Pinia 状态管理\n'
    '│   │   ├── api/               # API 调用封装\n'
    '│   │   └── router/            # 路由配置\n'
    '│   └── package.json\n'
    '├── README.md                   # 项目手册\n'
    '├── DEPLOY.md                   # 部署指南\n'
    '└── 项目功能说明书.docx          # 本文档'
)
p = doc.add_paragraph()
run = p.add_run(structure)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_heading('2.3 工作流程', level=2)

doc.add_paragraph(
    '当用户在对话框输入问题时，系统按以下流程处理：'
)

flow_steps = [
    '敏感词过滤 —— 检测并屏蔽违规内容',
    '保存消息 —— 用户消息写入数据库',
    '加载历史 —— 从数据库加载最近 10 轮对话作为上下文',
    '拒绝规则 —— 检查是否触发拒绝回答规则',
    '对话规则 —— 检查是否命中预设对话规则',
    '工具调用 —— 检查是否需要调用学生查询/请假等工具',
    '意图识别 —— 识别用户意图（问候/帮助/知识库等）',
    '问题重写 —— 语义纠错 + 代词消解 + 规则改写',
    'RAG 检索 —— 从多个知识库检索相关文档片段',
    'LLM 生成 —— 将检索结果 + 对话历史发送给 DeepSeek，生成回答',
    '后处理 —— 截断过长回答、敏感词过滤、生成推荐追问',
]
for i, step in enumerate(flow_steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

doc.add_page_break()

# ========== 三、功能模块详解 ==========
doc.add_heading('三、功能模块详解', level=1)

# 3.1
doc.add_heading('3.1 用户认证与权限管理', level=2)
doc.add_paragraph('系统支持完整的用户生命周期管理和基于角色的访问控制（RBAC）。')

doc.add_heading('功能列表', level=3)
auth_features = [
    '用户注册 —— 用户名 3-20 位字母数字，密码不少于 6 位',
    '用户登录 —— JWT Token 认证，5 次失败锁定 5 分钟',
    'Token 持久化 —— 存储在 localStorage，刷新页面自动恢复登录状态',
    '角色管理 —— 三种角色：admin（管理员）、kb_admin（知识库管理员）、user（普通用户）',
    '用户管理 —— 管理员可编辑、删除所有用户；普通用户只能编辑自己',
    '权限控制 —— 工具调用仅限 admin 和 kb_admin 角色',
    '密码安全 —— PBKDF2-SHA256 哈希（10 万轮迭代 + 随机盐值）',
]
for f in auth_features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('默认账号', level=3)
table = doc.add_table(rows=1, cols=3, style='Table Grid')
hdr = table.rows[0].cells
hdr[0].text = '用户名'
hdr[1].text = '密码'
hdr[2].text = '角色'
for u, p, r in [('admin', '123456', '管理员'), ('zzl', '123456', '普通用户'), ('yanhao', '123456', '普通用户'), ('YH666', '123456', '普通用户')]:
    row = table.add_row().cells
    row[0].text = u
    row[1].text = p
    row[2].text = r

doc.add_paragraph()

# 3.2
doc.add_heading('3.2 知识库管理', level=2)
doc.add_paragraph('知识库是 RAG 系统的核心组织单元，每个知识库对应一个 ChromaDB 集合。')

kb_features = [
    '创建知识库 —— 设置名称、描述、所属部门、负责人',
    '编辑知识库 —— 修改名称、描述、Embedding 模型',
    '删除知识库 —— 级联删除关联文档和向量数据',
    '克隆知识库 —— 复制配置和文档记录到新库（需重新向量化）',
    '筛选查询 —— 按名称、部门、时间范围筛选',
    'Embedding 模型 —— 支持 text2vec-base-chinese、bge-small/base/large-zh',
]
for f in kb_features:
    doc.add_paragraph(f, style='List Bullet')

# 3.3
doc.add_heading('3.3 文档管理', level=2)
doc.add_paragraph('支持文档的完整生命周期管理，从上传到向量化再到检索。')

doc_features = [
    '上传文档 —— 支持 PDF、Word、TXT、Markdown、图片、音频等 13 种格式',
    '批量上传 —— 一次选择多个文件同时上传',
    '文档解析 —— 自动解析文件内容，切分为语义片段',
    '向量化 —— 使用 Embedding 模型将文本片段转为向量存入 ChromaDB',
    '文档预览 —— 查看原始文件内容（前 5000 字符）',
    '拆分预览 —— 查看文本切分后的片段效果',
    '向量预览 —— 查看向量维度、相似度得分等调试信息',
    '文档编辑 —— 修改文件内容后自动重新向量化',
    '文档版本 —— 记录修改历史，支持回滚到历史版本',
    '文档删除 —— 删除文档同时清理 ChromaDB 向量数据',
    '向量去重 —— 检测并删除重复向量，优化检索效果',
    '上传限制 —— 单文件最大 50MB，文件类型白名单校验',
]
for f in doc_features:
    doc.add_paragraph(f, style='List Bullet')

# 3.4
doc.add_heading('3.4 AI 对话系统', level=2)
doc.add_paragraph('对话系统是面向终端用户的核心交互界面。')

chat_features = [
    '多轮对话 —— 维护最近 10 轮上下文，支持追问理解',
    '多知识库选择 —— 对话时可选择检索一个或多个知识库',
    '快捷提问 —— 预设常见问题模板，一键发送',
    '会话管理 —— 创建、切换、删除对话会话',
    '会话级参数 —— 每个对话可独立设置 top_k、min_score',
    '回答长度限制 —— 可配置最大回答字数',
    '推荐追问 —— AI 回答后自动生成 3 个相关追问建议',
    '会话数据隔离 —— 退出登录时清空会话，登录时加载对应用户的历史',
]
for f in chat_features:
    doc.add_paragraph(f, style='List Bullet')

# 3.5
doc.add_heading('3.5 RAG 检索引擎', level=2)
doc.add_paragraph('RAG（检索增强生成）引擎负责从知识库中检索相关信息，为 LLM 提供上下文。')

rag_features = [
    '文档解析 —— 支持 PDF（pdfplumber）、Word（python-docx）、纯文本等格式',
    '文本切分 —— 可配置 chunk_size（切片大小）和 overlap（重叠长度）',
    '向量化 —— sentence-transformers 生成文本向量',
    '多库检索 —— 同时检索多个知识库，按相似度排序取 top_k',
    '相似度过滤 —— 设置 min_score 阈值，过滤低相关性结果',
    '向量去重 —— 自动检测并删除重复向量',
]
for f in rag_features:
    doc.add_paragraph(f, style='List Bullet')

# 3.6
doc.add_heading('3.6 工具调用系统', level=2)
doc.add_paragraph('工具调用系统允许 AI 通过关键词匹配自动执行预定义的结构化操作。')

tool_features = [
    '查询学生 —— 输入"查询学生 张三"或"查询学生 学号2023001"',
    '请假申请 —— 输入"帮我请假 学号2023001 2024-01-01 到 2024-01-05 因为生病"',
    '请假状态 —— 输入"请假状态 学号2023001"查看请假记录',
    '权限校验 —— 仅 admin 和 kb_admin 角色可调用工具',
    '友好回复 —— 工具结果转换为自然语言返回，不显示 JSON',
]
for f in tool_features:
    doc.add_paragraph(f, style='List Bullet')

# 3.7
doc.add_heading('3.7 日志与数据分析', level=2)

log_features = [
    '对话日志 —— 记录所有对话消息，支持按用户、会话、时间筛选',
    '会话列表 —— 按会话维度展示对话摘要',
    '数据统计 —— 对话总数、用户数、平均对话轮次',
    '数据可视化 —— 热门问题、检索成功率、用户活跃度图表',
    'CSV 导出 —— 一键导出对话日志为 CSV 文件',
    '过期清理 —— 按天数自动清理历史日志',
]
for f in log_features:
    doc.add_paragraph(f, style='List Bullet')

# 3.8
doc.add_heading('3.8 系统配置', level=2)

config_features = [
    'chunk_size —— 文本切片大小（默认 500 字符）',
    'overlap —— 切片重叠长度（默认 50 字符）',
    'top_k —— 检索返回的结果数量（默认 5）',
    'min_score —— 相似度阈值（默认 0.3）',
    'llm_model —— 使用的 LLM 模型名称',
    'max_context_turns —— 多轮对话上下文轮数（默认 10）',
    'max_answer_length —— 回答最大字数（默认 2000）',
    'quick_questions —— 快捷提问列表',
]
for f in config_features:
    doc.add_paragraph(f, style='List Bullet')

# 3.9
doc.add_heading('3.9 多模态输入', level=2)

mm_features = [
    '图片 OCR —— 上传图片后使用 PaddleOCR 识别文字内容',
    '文件上传 —— 支持 PDF、Word、TXT 等格式，自动提取文本',
    '文件标签 —— 上传后显示文件名标签，用户可追加补充文字',
    '内容合并 —— OCR 提取的文字与用户补充文字合并后发送给 AI',
]
for f in mm_features:
    doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()

# ========== 四、部署指南 ==========
doc.add_heading('四、部署指南', level=1)

doc.add_heading('4.1 环境要求', level=2)

table = doc.add_table(rows=1, cols=3, style='Table Grid')
hdr = table.rows[0].cells
hdr[0].text = '组件'
hdr[1].text = '版本要求'
hdr[2].text = '说明'
env_data = [
    ('Python', '>= 3.10', '后端运行环境'),
    ('Node.js', '>= 18', '前端构建工具'),
    ('DeepSeek API Key', '-', '对话生成服务（需注册获取）'),
    ('Redis', '>= 6.0', '可选，缓存加速（未安装时自动使用内存缓存）'),
]
for comp, ver, desc in env_data:
    row = table.add_row().cells
    row[0].text = comp
    row[1].text = ver
    row[2].text = desc

doc.add_paragraph()

doc.add_heading('4.2 后端部署', level=2)

backend_steps = [
    '克隆项目：git clone <仓库地址> smart_campus_agent',
    '进入后端目录：cd smart_campus_agent/backend',
    '创建虚拟环境：python -m venv venv → venv\\Scripts\\activate',
    '安装依赖：pip install -r requirements.txt',
    '配置环境变量：设置系统环境变量 DeepSeek_API_KEY（值为你的 API 密钥）',
    '启动后端：python -m uvicorn app.main:app --reload --port 8000',
    '验证：浏览器打开 http://localhost:8000/docs 查看 API 文档',
]
for i, step in enumerate(backend_steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

doc.add_heading('4.3 前端部署', level=2)

frontend_steps = [
    '进入前端目录：cd smart_campus_agent/frontend',
    '安装依赖：npm install',
    '启动开发服务器：npm run dev',
    '访问：浏览器打开 http://localhost:5173',
]
for i, step in enumerate(frontend_steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

doc.add_heading('4.4 一键启动脚本', level=2)

doc.add_paragraph('项目提供了 Python 版一键启动脚本，位于 E:\\实训\\Yan\\ 目录：')

script_data = [
    ('启动.py', '同时启动后端和前端服务，各自在独立窗口运行'),
    ('关闭.py', '关闭所有服务（uvicorn + vite + node 进程）'),
    ('重启.py', '先关闭再启动，等同于一键重启'),
]
for name, desc in script_data:
    doc.add_paragraph(f'{name} —— {desc}', style='List Bullet')

doc.add_page_break()

# ========== 五、使用指南 ==========
doc.add_heading('五、使用指南', level=1)

doc.add_heading('5.1 登录系统', level=2)
doc.add_paragraph('打开浏览器访问 http://localhost:5173，进入登录页面。输入用户名和密码（默认 admin/123456），点击登录。登录成功后进入对话页面。')

doc.add_heading('5.2 知识库操作', level=2)
doc.add_paragraph('点击左侧导航栏底部的"知识库管理"按钮，进入知识库管理页面。')

kb_steps = [
    '创建知识库：点击"新建知识库"按钮，填写名称、描述、所属部门等信息',
    '上传文档：进入文档管理页面，点击"选择文件"上传 PDF/Word/TXT 等文档',
    '等待解析：文档上传后自动解析、切分、向量化，进度条显示处理进度',
    '开始对话：回到对话页面，选择对应的知识库，输入问题即可获得基于文档的回答',
]
for i, step in enumerate(kb_steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

doc.add_heading('5.3 智能对话', level=2)
doc.add_paragraph('在对话页面中，你可以：')

chat_guide = [
    '直接输入问题，AI 会基于知识库内容回答',
    '上传文件后输入补充说明，点击发送（文件内容会自动 OCR 识别后合并发送）',
    '点击快捷提问按钮快速发送常见问题',
    '在知识库下拉框中切换不同的知识库',
    '点击左侧会话列表切换或创建新对话',
]
for g in chat_guide:
    doc.add_paragraph(g, style='List Bullet')

doc.add_heading('5.4 工具调用示例', level=2)

table = doc.add_table(rows=1, cols=2, style='Table Grid')
hdr = table.rows[0].cells
hdr[0].text = '输入内容'
hdr[1].text = '预期结果'
tool_examples = [
    ('我要请假', '提示"请提供学号，例如：帮我请假 学号2023001 2024-01-01 到 2024-01-05"'),
    ('帮我请假 学号2023001 2024-01-01 到 2024-01-05 因为生病', '返回"请假申请已提交！工单号：X，状态：pending"'),
    ('请假状态 学号2023001', '返回该学生的所有请假记录列表'),
    ('查询学生 学号2023001', '返回学生信息：姓名、学号、班级、电话'),
    ('请假结束后如何取消', 'AI 基于知识库/LLM 回答（不触发工具）'),
]
for inp, out in tool_examples:
    row = table.add_row().cells
    row[0].text = inp
    row[1].text = out

doc.add_paragraph()

doc.add_heading('5.5 日志查看', level=2)
doc.add_paragraph('管理员可通过 API 查看系统日志和数据统计：')

log_guide = [
    'GET /api/v1/logs/ —— 查看对话日志列表',
    'GET /api/v1/logs/stats —— 查看统计数据',
    'GET /api/v1/logs/analytics —— 查看数据分析',
    'GET /api/v1/logs/export —— 导出 CSV 格式日志文件',
]
for g in log_guide:
    doc.add_paragraph(g, style='List Bullet')

doc.add_page_break()

# ========== 六、API 接口说明 ==========
doc.add_heading('六、API 接口说明', level=1)
doc.add_paragraph('所有 API 接口的基础路径为 http://localhost:8000/api/v1/，需要 JWT Token 认证（除注册和登录外）。')

doc.add_paragraph('访问 http://localhost:8000/docs 可查看完整的交互式 API 文档（Swagger UI）。')

table = doc.add_table(rows=1, cols=4, style='Table Grid')
hdr = table.rows[0].cells
hdr[0].text = '模块'
hdr[1].text = '方法'
hdr[2].text = '路径'
hdr[3].text = '说明'

api_data = [
    ('认证', 'POST', '/auth/register', '用户注册'),
    ('认证', 'POST', '/auth/login', '用户登录'),
    ('知识库', 'GET', '/kb/', '知识库列表'),
    ('知识库', 'POST', '/kb/', '创建知识库'),
    ('知识库', 'PUT', '/kb/{id}', '编辑知识库'),
    ('知识库', 'DELETE', '/kb/{id}', '删除知识库'),
    ('知识库', 'POST', '/kb/{id}/clone', '克隆知识库'),
    ('文档', 'POST', '/doc/upload', '上传文档'),
    ('文档', 'GET', '/doc/{kb_id}', '文档列表'),
    ('文档', 'DELETE', '/doc/{id}', '删除文档'),
    ('文档', 'GET', '/doc/{id}/preview', '预览文档'),
    ('文档', 'PUT', '/doc/{id}/edit', '编辑文档'),
    ('文档', 'POST', '/doc/multimodal', '多模态上传'),
    ('文档', 'GET', '/doc/{id}/vectors', '向量预览'),
    ('对话', 'POST', '/chat/', '发送消息'),
    ('对话', 'GET', '/chat/{session_id}', '获取历史'),
    ('对话', 'DELETE', '/chat/{session_id}', '删除会话'),
    ('日志', 'GET', '/logs/', '日志列表'),
    ('日志', 'GET', '/logs/stats', '数据统计'),
    ('日志', 'GET', '/logs/export', 'CSV 导出'),
    ('配置', 'GET', '/settings/', '获取配置'),
    ('配置', 'PUT', '/settings/', '更新配置'),
]
for mod, method, path, desc in api_data:
    row = table.add_row().cells
    row[0].text = mod
    row[1].text = method
    row[2].text = path
    row[3].text = desc

doc.add_paragraph()

doc.add_heading('Python 调用示例', level=2)

code = """import requests

# 1. 登录获取 Token
resp = requests.post("http://localhost:8000/api/v1/auth/login",
    json={"username": "admin", "password": "123456"})
token = resp.json()["access_token"]

# 2. 发送对话消息
headers = {"Authorization": f"Bearer {token}"}
resp = requests.post("http://localhost:8000/api/v1/chat/",
    json={"session_id": "test", "kb_id": 1, "query": "你好"},
    headers=headers)
print(resp.json()["answer"])"""

p = doc.add_paragraph()
run = p.add_run(code)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_page_break()

# ========== 七、常见问题 ==========
doc.add_heading('七、常见问题', level=1)

faqs = [
    ('启动后前端页面空白', '检查后端是否启动，确认 vite.config.js 中的代理配置指向 localhost:8000'),
    ('对话返回"知识库中没有找到信息"', '确保已上传文档并等待向量化完成（进度 100%），且在对话时选择了正确的知识库'),
    ('DeepSeek API 调用失败', '检查环境变量 DeepSeek_API_KEY 是否正确设置，API 余额是否充足'),
    ('ChromaDB 首次加载慢', '首次使用会下载 Embedding 模型（约 80MB），之后会缓存'),
    ('pip install 报错 SSL', '执行 $env:NO_PROXY="*"; $env:no_proxy="*" 绕过代理'),
    ('pydantic 版本冲突', '执行 pip install pydantic==2.9.2 pydantic-core==2.23.4 pydantic-settings'),
    ('前端跨域错误', '确认后端 CORS 配置包含前端地址（默认允许 localhost:5173）'),
    ('工具调用返回"未找到该学生"', '确认学生表中有对应学号的数据，可通过 SQL 插入测试数据'),
]

for q, a in faqs:
    p = doc.add_paragraph()
    run = p.add_run(f'Q: {q}')
    run.font.bold = True
    doc.add_paragraph(f'A: {a}')

# ========== 保存 ==========
output_path = r"E:\实训资料\Smart_Campus_Agent_项目文档.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
